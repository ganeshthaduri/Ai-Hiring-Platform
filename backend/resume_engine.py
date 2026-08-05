"""
resume_engine.py
================
The analysis engine behind the Resume Readiness Check tool.

WHAT THIS FILE DOES
-----------------------------------------------------------------
1. FILE PARSING            -> pulls raw text out of an uploaded .pdf / .docx resume
2. ENTITY EXTRACTION       -> spaCy (en_core_web_sm) + heuristics for name/email/phone
3. SECTION SEGMENTATION    -> splits the resume into labeled sections
4. SKILL CATEGORIZATION    -> skills grouped into Languages / Frameworks / Libraries /
                              Databases / Cloud & DevOps / Tools & Platforms / Soft Skills
5. SEMANTIC SKILL MATCHING -> a JD skill can match a resume even if the exact
                              string isn't there, via synonym resolution and
                              (when available) sentence-embedding similarity --
                              this replaces plain substring "AI" matching
6. SEMANTIC SIMILARITY     -> the headline resume<->JD match score. Uses
                              sentence-transformers embeddings when available;
                              falls back to TF-IDF cosine similarity and is
                              labeled as such so the UI never overstates what
                              produced the number
7. RESUME COMPLETENESS     -> did the resume include the sections/contact info
                              an ATS or recruiter expects?
8. RESUME STRUCTURE        -> is the resume formatted in a way a parser can
                              read reliably (clear headers, bullets, length)?
9. PROJECT ANALYSIS        -> are projects relevant, quantified, and using
                              strong action language?
10. INTERVIEW READINESS    -> a practical readiness read combining the above
11. EXPLAINABILITY         -> every score returned includes the evidence used
                              to compute it, and a to_dict() for downstream use

Run this file directly to see it work end-to-end on sample data:
    python3 resume_engine.py
"""
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))

import re
import json
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional

import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from backend.skills_db import SKILLS_DB, SOFT_SKILLS, SKILL_SYNONYMS, all_categories

# ---------------------------------------------------------------------------
# 1. LOAD THE PRETRAINED spaCy MODEL ONCE
# ---------------------------------------------------------------------------
try:
    NLP = spacy.load("en_core_web_sm")
except OSError as e:
    raise OSError(
        "The pretrained spaCy model 'en_core_web_sm' isn't installed yet.\n"
        "Run this once, then try again:\n\n"
        "    python -m spacy download en_core_web_sm\n"
    ) from e


# ---------------------------------------------------------------------------
# 2. FILE -> TEXT
# ---------------------------------------------------------------------------
def extract_text_from_pdf(path: str) -> str:
    """
    FIX: pdfplumber's extract_text() decides where to insert a space between
    two words based on the horizontal gap between their characters, using a
    default x_tolerance of 3pt. Several common resume templates (Word/Canva
    exports especially) render text with tighter character/word spacing than
    that threshold, so pdfplumber sees no gap at all and glues every word in
    the document together -- "Professional Summary" becomes
    "ProfessionalSummary", "Technical Skills" becomes "TechnicalSkills", etc.

    That silently breaks every downstream step that matches on word
    boundaries: section header detection (SECTION_HEADERS), skill surface
    forms, degree/email/phone regexes -- even though the resume visibly has
    all of that content. Lowering x_tolerance makes pdfplumber sensitive to
    the smaller-but-real gaps these templates use, without merging unrelated
    text. We try progressively smaller values and keep the first one that
    actually yields space-separated words.
    """
    import pdfplumber
    import re as _re

    def _looks_glued(t: str) -> bool:
        # Heuristic: when spaces between words are dropped, whole runs of
        # real words get concatenated into single long tokens (e.g.
        # "ProfessionalSummary", "DataProfessionalwithastrong..."), which
        # pulls the average "word" length way up. Normal English prose
        # (including resumes full of jargon/acronyms) averages ~4-7 chars
        # per word; glued text runs well into the teens.
        words = _re.findall(r"[A-Za-z]+", t)
        if len(words) < 20:
            return False
        avg_len = sum(len(w) for w in words) / len(words)
        return avg_len > 9

    def _extract_page(page):
        default = page.extract_text() or ""
        if not _looks_glued(default):
            return default
        for tol in (2, 1, 0.5):
            candidate = page.extract_text(x_tolerance=tol) or ""
            if not _looks_glued(candidate):
                return candidate
        return default

    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text.append(_extract_page(page))
    return "\n".join(text)


def extract_text_from_docx(path: str) -> str:
    """
    Pulls text out of a .docx file IN DOCUMENT ORDER, including text that
    lives inside tables.

    FIX: python-docx's `doc.paragraphs` only walks top-level body paragraphs.
    A lot of resume templates lay out the header, skills, or a two-column
    layout using TABLES (or the whole resume is one big table). Reading only
    `doc.paragraphs` silently drops all of that text -- which then makes
    perfectly real sections (e.g. "Skills") look completely absent to the
    rest of the pipeline, even though the resume clearly has them. Walking
    both paragraphs AND tables (recursively, since tables can be nested)
    fixes that.
    """
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    doc = docx.Document(path)

    def iter_block_items(parent):
        if hasattr(parent, "element"):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._tc
        for child in parent_elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    def walk(parent, parts):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                t = block.text.strip()
                if t:
                    parts.append(t)
            elif isinstance(block, Table):
                for row in block.rows:
                    row_texts = []
                    for cell in row.cells:
                        # Recurse into the cell in case it contains nested
                        # tables/paragraphs, and also grab plain cell.text
                        # as a safety net for any content the recursion misses.
                        cell_parts = []
                        walk(cell, cell_parts)
                        cell_text = "\n".join(cell_parts).strip() or cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        parts.append(" | ".join(row_texts))
        return parts

    parts = walk(doc, [])
    text = "\n".join(parts)

    # Safety net: if the structured walk somehow produced nothing (unusual
    # docx internals), fall back to the simple paragraph-only read plus a
    # raw table dump so we never return an empty string for a non-empty file.
    if not text.strip():
        fallback_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        fallback_parts.append(cell.text.strip())
        text = "\n".join(fallback_parts)

    return text


def load_document(path: str) -> str:
    if path.lower().endswith(".pdf"):
        return extract_text_from_pdf(path)
    if path.lower().endswith(".docx"):
        return extract_text_from_docx(path)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# ---------------------------------------------------------------------------
# 3. SECTION SEGMENTATION (Education / Experience / Projects / Certifications / Skills)
# ---------------------------------------------------------------------------
SECTION_HEADERS = {
    "summary":        ["summary","Summary", "professional summary","Professional Summary", "objective", "profile", "about me",
                        "career objective", "personal summary"],
    "education":      ["education", "academic background","Education", "academic qualifications", "qualifications",
                        "educational qualifications", "academic details"],
    "experience":     ["experience", "work experience", "professional experience",
                        "employment history", "work history","Internship EXperience","INTERNSHIP EXPERIENCE""Experience", "internship experience",
                        "internships", "internship"],
    "projects":       ["projects","Projects", "academic projects", "personal projects", "key projects",
                        "project details", "project work"],
    "certifications": ["certifications","Certifications","Certificates" "certificates", "licenses", "certifications & licenses",
                        "licenses & certifications", "certifications and licenses", "courses"],
    "skills":         ["skills","Skills","Technical Skills", "technical skills", "core competencies", "skills & tools",
                        "key skills", "skills and tools", "skill set", "areas of expertise",
                        "technical skill set", "skills summary"],
}

# Characters that commonly decorate a header line (bullets, markdown bold/
# underline markers, box-drawing dashes, colons used as a header/ content
# separator, etc.) -- stripped before we try to match against SECTION_HEADERS.
_HEADER_DECOR_CHARS = r"\-\*_=#•◦▪·»«~^><\|"


def _normalize_header_text(raw: str) -> str:
    """Collapses a chunk of text down to a bare lowercase phrase so headers
    formatted as '**Skills**', '== SKILLS ==', '• Skills', 'Skills:' etc.
    all normalize the same way."""
    t = raw.strip()
    t = re.sub(rf"^[{_HEADER_DECOR_CHARS}\s]+", "", t)
    t = re.sub(rf"[{_HEADER_DECOR_CHARS}\s]+$", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip().lower()


def _match_section_key(clean: str) -> Optional[str]:
    """Returns the section key if `clean` looks like one of our known
    headers -- exact match, or the header phrase plus a couple of harmless
    extra words (e.g. 'technical skills required', 'skills & tools used')."""
    if not clean:
        return None
    for section_key, keywords in SECTION_HEADERS.items():
        for kw in keywords:
            if clean == kw:
                return section_key
            if clean.startswith(kw + " ") and len(clean) <= len(kw) + 20:
                return section_key
    return None


def segment_resume_sections(resume_text: str) -> Dict[str, str]:
    """
    Splits raw resume text into labeled sections by detecting header lines.
    Sections that couldn't be detected are simply absent from the dict --
    the UI is responsible for showing a "not clearly detected" message.

    Handles two header styles:
      1. Header alone on its own line ("Skills", "**Skills**", "== SKILLS ==")
      2. Header + inline content on the same line ("Skills: Python, SQL, ...")
         -- previously this second style was silently ignored whenever the
         combined line was longer than 40 characters, which made an entire,
         genuinely-present section (most commonly Skills) look completely
         missing to the rest of the pipeline.
    """
    lines = resume_text.split("\n")
    header_at_line: Dict[int, str] = {}
    inline_content: Dict[int, str] = {}

    for i, raw_line in enumerate(lines):
        stripped = raw_line.strip()
        if not stripped:
            continue

        # --- Case 1: header alone on its own line ---
        clean = _normalize_header_text(stripped)
        if clean and len(clean) <= 40:
            key = _match_section_key(clean)
            if key:
                header_at_line[i] = key
                continue

        # --- Case 2: "Header: content" or "Header - content" on one line ---
        m = re.match(r"^([A-Za-z][A-Za-z &/]{1,30}?)\s*[:\-–]\s*(.+)$", stripped)
        if m:
            head_clean = _normalize_header_text(m.group(1))
            key = _match_section_key(head_clean)
            if key:
                header_at_line[i] = key
                inline_content[i] = m.group(2).strip()

    if not header_at_line:
        return {}

    ordered_positions = sorted(header_at_line.keys())
    sections: Dict[str, str] = {}
    for pos_idx, line_idx in enumerate(ordered_positions):
        key = header_at_line[line_idx]
        start = line_idx + 1
        end = ordered_positions[pos_idx + 1] if pos_idx + 1 < len(ordered_positions) else len(lines)
        body_lines = lines[start:end]
        content_parts = []
        if line_idx in inline_content:
            content_parts.append(inline_content[line_idx])
        content_parts.append("\n".join(body_lines).strip())
        content = "\n".join(p for p in content_parts if p).strip()
        if content:
            sections[key] = (sections[key] + "\n" + content) if key in sections else content

    return sections


# ---------------------------------------------------------------------------
# 4. STRUCTURED FIELD EXTRACTION (candidate summary card)
# ---------------------------------------------------------------------------
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[ -]?)?\d{10}|(\+?\d{1,3}[ -]?)?\d{3,5}[ -]\d{3,5}[ -]?\d{0,5}")
EXPERIENCE_RE = re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)", re.IGNORECASE)


def extract_candidate_name(resume_text: str) -> str:
    """
    Extracts the candidate's name using a hybrid approach: heuristics on the
    top section of the resume, cross-checked against spaCy NER.
    """
    lines = resume_text.split("\n")
    top_lines = []

    stop_headers = {
        "summary", "education", "experience", "projects", "certifications", "skills",
        "objective", "profile", "about me", "work history", "employment", "academic",
        "contact", "contact info", "contact information", "publications", "interests",
        "hobbies", "languages", "achievements", "awards", "declarations", "declaration"
    }

    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
        normalized = re.sub(r'[^a-z0-9\s]', '', cleaned_line.lower()).strip()
        if normalized in stop_headers:
            break
        top_lines.append(cleaned_line)
        if len(top_lines) >= 8:
            break

    excluded_terms = {
        "resume", "cv", "curriculum", "vitae", "portfolio", "profile", "contact",
        "page", "address", "phone", "email", "mobile", "github", "linkedin",
        "website", "job", "applicant", "aspirant", "candidate", "engineer",
        "developer", "analyst", "manager", "architect", "consultant", "student",
        "intern", "designer", "coder", "programmer", "lead", "head", "director", "vp",
        "python", "java", "pandas", "numpy", "django", "fastapi", "sql", "git",
        "aws", "kubernetes", "c++", "c#", "html", "css", "javascript", "react",
        "angular", "node", "docker", "gcp", "azure", "ml", "machine", "learning",
        "data", "scientist", "software", "certified", "professional", "summary",
        "hiring", "talent", "recruitment", "recruiter", "phone number", "present"
    }

    degree_pattern = re.compile(r",?\s*\b(phd|ph\.d|mtech|btech|m\.tech|b\.tech|ms|bs|mba|pmp|m\.d|md|b\.sc|m\.sc)\b", re.IGNORECASE)
    prefix_pattern = re.compile(r"\b(dr|mr|ms|mrs|prof)\.?\s+", re.IGNORECASE)

    candidates = []

    for line in top_lines:
        segments = re.split(r'\s+[-|/•·]\s+|\s*,\s*', line)
        first_segment = segments[0].strip()

        if not first_segment:
            continue
        if not first_segment[0].isalpha():
            continue
        if "@" in first_segment:
            continue
        if any(term in first_segment.lower() for term in ["http:", "https:", ".com", ".org", ".edu", ".net", ".io", "github.com", "linkedin.com"]):
            continue
        digits_count = sum(c.isdigit() for c in first_segment)
        if digits_count > 3:
            continue

        cleaned = degree_pattern.sub("", first_segment).strip()
        cleaned = prefix_pattern.sub("", cleaned).strip()
        cleaned = cleaned.rstrip(",;|").strip()

        if not cleaned:
            continue

        words = cleaned.split()
        if len(words) < 2 or len(words) > 5:
            continue
        if any(w.lower() in excluded_terms for w in words):
            continue

        is_capitalized = True
        for w in words:
            w_clean = re.sub(r'[^a-zA-Z]', '', w)
            if not w_clean:
                is_capitalized = False
                break
            if not (w_clean[0].isupper() or w_clean.isupper()):
                is_capitalized = False
                break
        if is_capitalized:
            candidates.append(cleaned)

    spacy_names = []
    doc = NLP(resume_text[:2000])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            ent_clean = degree_pattern.sub("", ent.text).strip()
            ent_clean = prefix_pattern.sub("", ent_clean).strip()
            ent_words = ent_clean.split()
            if any(w.lower() in excluded_terms for w in ent_words):
                continue
            if "@" in ent_clean or sum(c.isdigit() for c in ent_clean) > 2:
                continue
            if not ent_clean or not ent_clean[0].isalpha():
                continue
            spacy_names.append(ent_clean)

    for s_name in spacy_names:
        for cand in candidates:
            if s_name.lower() in cand.lower() or cand.lower() in s_name.lower():
                return cand

    if candidates:
        return candidates[0]

    if spacy_names:
        return spacy_names[0]

    for line in top_lines:
        clean_line = line.strip()
        if not clean_line:
            continue
        if "@" in clean_line:
            continue
        if sum(c.isdigit() for c in clean_line) > 3:
            continue
        segments = re.split(r'\s+[-|/•·]\s+|\s*,\s*', clean_line)
        candidate = segments[0].strip()
        candidate = degree_pattern.sub("", candidate).strip()
        candidate = prefix_pattern.sub("", candidate).strip()
        if candidate and candidate[0].isalpha():
            if candidate.islower():
                return candidate.title()
            return candidate

    return "Unknown Candidate"


import re

def extract_candidate_profile(resume_text: str) -> Dict:
    name = extract_candidate_name(resume_text)

    email_match = EMAIL_RE.search(resume_text)
    phone_match = PHONE_RE.search(resume_text)
    years_match = EXPERIENCE_RE.search(resume_text)

    lowered = resume_text.lower()

    # ---------------- Degree ---------------- #
    degree_patterns = {
        r"\b(ph\.?d|doctor of philosophy)\b": "PhD",

        r"\b(master(?:'s)?|masters|m\.?tech|mtech|m\.?e|me|mba|mca|m\.?sc|msc)\b":
            "Master",

        r"\b(bachelor(?:'s)?|bachelors|b\.?tech|btech|b\.?e|be|b\.?sc|bsc|bca|bba|b\.?com|bcom)\b":
            "Bachelor",

        r"\b(diploma|polytechnic)\b":
            "Diploma"
    }

    highest_degree = "Not specified"

    for pattern, label in degree_patterns.items():
        if re.search(pattern, lowered):
            highest_degree = label
            break

    # ---------------- Field ---------------- #
    field_patterns = {
        r"\bcomputer science and engineering\b": "Computer Science",
        r"\bcomputer science\b": "Computer Science",
        r"\bcse\b": "Computer Science",

        r"\binformation technology\b": "Information Technology",
        r"\bit\b": "Information Technology",

        r"\bsoftware engineering\b": "Software Engineering",

        r"\bartificial intelligence\b": "Artificial Intelligence",

        r"\bmachine learning\b": "Machine Learning",

        r"\bdata science\b": "Data Science",

        r"\bstatistics\b": "Statistics",

        r"\bmathematics\b": "Mathematics",

        r"\beconomics\b": "Economics",

        r"\belectronics and communication\b": "Electronics",

        r"\bece\b": "Electronics",

        r"\bmechanical\b": "Mechanical Engineering",

        r"\bcivil\b": "Civil Engineering",

        r"\belectrical\b": "Electrical Engineering"
    }

    education_field = "Not specified"

    for pattern, label in field_patterns.items():
        if re.search(pattern, lowered):
            education_field = label
            break

    return {
        "name": name,
        "email": email_match.group(0) if email_match else "Not found",
        "phone": phone_match.group(0).strip() if phone_match else "Not found",
        "years_experience": float(years_match.group(1)) if years_match else None,
        "highest_education": highest_degree,
        "education_field": education_field,
    }


# ---------------------------------------------------------------------------
# 5. SKILL DETECTION + CATEGORIZATION (single source of truth -- used
#    everywhere so the UI never renders the same skill list twice under
#    two different labels)
# ---------------------------------------------------------------------------
def _compile_surface_form_pattern(form: str) -> re.Pattern:
    """Word-boundary match for alphanumeric surface forms so 'scala' doesn't
    fire inside 'scalable', 'go' doesn't fire inside 'good', etc. Forms with
    symbols (c++, c#, .net, ci/cd) can't use \\b cleanly, so those fall back
    to a plain (but still whitespace-aware) substring check."""
    if re.match(r"^[a-z0-9][a-z0-9 .\-]*[a-z0-9]$", form) or re.match(r"^[a-z0-9]$", form):
        return re.compile(r"(?<![a-z0-9])" + re.escape(form.strip()) + r"(?![a-z0-9])")
    return None


_SURFACE_FORM_PATTERNS: Dict[str, List] = {}
for _skill, _meta in SKILLS_DB.items():
    _compiled = []
    for _form in _meta["surface_forms"]:
        _pattern = _compile_surface_form_pattern(_form.strip())
        _compiled.append(_pattern if _pattern is not None else _form.strip())
    _SURFACE_FORM_PATTERNS[_skill] = _compiled


def find_skills(text: str) -> List[str]:
    """Exact/alias surface-form detection against the controlled taxonomy,
    using word-boundary matching to avoid false positives like 'scala'
    inside 'scalable' or 'go' inside 'good'."""
    lowered = text.lower()
    found = []
    for skill, patterns in _SURFACE_FORM_PATTERNS.items():
        for p in patterns:
            hit = p.search(lowered) if isinstance(p, re.Pattern) else (p in lowered)
            if hit:
                found.append(skill)
                break
    return found


def categorize_skills(skill_keys: List[str]) -> Dict[str, List[str]]:
    """Groups a flat list of skill keys into the taxonomy's categories.
    This is the ONE place skills get grouped for display -- callers should
    render this, not re-derive technical/soft lists separately."""
    grouped: Dict[str, List[str]] = {cat: [] for cat in all_categories()}
    for skill in skill_keys:
        cat = SKILLS_DB.get(skill, {}).get("category", "Other")
        grouped.setdefault(cat, []).append(skill)
    return {cat: sorted(set(skills)) for cat, skills in grouped.items() if skills}


def classify_resume_skills(resume_text: str) -> Dict[str, List[str]]:
    """All skills the resume mentions, categorized -- independent of any JD,
    so a candidate can see this even before pasting a job description."""
    return categorize_skills(find_skills(resume_text))


# ---------------------------------------------------------------------------
# 6. SEMANTIC SKILL MATCHING
#    A JD skill counts as covered if:
#      a) it's an exact/alias surface-form match in the resume (find_skills), or
#      b) a known synonym/abbreviation resolves it to a resume skill, or
#      c) (when an embeddings backend is available) the JD skill phrase is
#         semantically close to text the resume actually contains
#    Each match records HOW it was matched, so the UI can be honest about it
#    instead of presenting every match as identical "AI matched" chips.
#    This is the ONE engine used everywhere skills get matched (global chips
#    AND the section-wise Technical/Soft Skills tabs), so scoring reflects
#    MEANING, not just literal keyword overlap.
# ---------------------------------------------------------------------------
_EMBED_MODEL = None
_EMBED_MODEL_LOAD_ATTEMPTED = False


def _get_embedding_model():
    """Lazily loads a sentence-embedding model if the package + weights are
    available. Returns None (rather than raising) if it can't be loaded, so
    the rest of the app degrades gracefully instead of crashing."""
    global _EMBED_MODEL, _EMBED_MODEL_LOAD_ATTEMPTED
    if _EMBED_MODEL_LOAD_ATTEMPTED:
        return _EMBED_MODEL
    _EMBED_MODEL_LOAD_ATTEMPTED = True
    try:
        from sentence_transformers import SentenceTransformer
        _EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        _EMBED_MODEL = None
    return _EMBED_MODEL


def _embedding_similarity(a: str, b: str) -> Optional[float]:
    model = _get_embedding_model()
    if model is None:
        return None
    try:
        emb = model.encode([a, b])
        sim = cosine_similarity([emb[0]], [emb[1]])[0][0]
        return float(sim)
    except Exception:
        return None


SEMANTIC_MATCH_THRESHOLD = 0.55


def semantic_skill_match(resume_text: str, resume_skills: List[str], jd_skills: List[str]) -> Dict:
    """
    Returns matched/missing JD skills, where "matched" includes exact,
    synonym-resolved, and (if available) embedding-based semantic matches.
    """
    resume_skill_set = set(resume_skills)
    resume_lower = resume_text.lower()

    matched = []   # list of {"skill":..., "match_type":..., "score": float|None}
    missing = []

    for skill in jd_skills:
        if skill in resume_skill_set:
            matched.append({"skill": skill, "match_type": "exact", "score": 1.0})
            continue

        # synonym resolution: does a synonym of this skill (or a synonym
        # pointing TO this skill) appear as a resume skill?
        synonym_hit = None
        for alias, canonical in SKILL_SYNONYMS.items():
            if canonical == skill and alias in resume_lower:
                synonym_hit = alias
                break
        if synonym_hit:
            matched.append({"skill": skill, "match_type": "synonym", "score": 0.9, "matched_via": synonym_hit})
            continue

        # embedding-based semantic match against the resume's own skill list
        best_score = None
        best_via = None
        if resume_skills:
            for r_skill in resume_skills:
                sim = _embedding_similarity(skill, r_skill)
                if sim is not None and (best_score is None or sim > best_score):
                    best_score = sim
                    best_via = r_skill
        if best_score is not None and best_score >= SEMANTIC_MATCH_THRESHOLD:
            matched.append({"skill": skill, "match_type": "semantic", "score": round(best_score, 3), "matched_via": best_via})
            continue

        missing.append(skill)

    technical_matched = [m for m in matched if m["skill"] not in SOFT_SKILLS]
    technical_missing = [s for s in missing if s not in SOFT_SKILLS]
    soft_matched = [m for m in matched if m["skill"] in SOFT_SKILLS]
    soft_missing = [s for s in missing if s in SOFT_SKILLS]

    return {
        "matched": matched,
        "missing": missing,
        "technical": {"matched": technical_matched, "missing": technical_missing},
        "soft": {"matched": soft_matched, "missing": soft_missing},
        "used_embeddings": _get_embedding_model() is not None,
    }


# ---------------------------------------------------------------------------
# 7. RESUME <-> JD SEMANTIC SIMILARITY (the headline match score)
#    Prefers real sentence embeddings; falls back to TF-IDF cosine similarity
#    when embeddings aren't available, and is transparent about which one
#    produced the number so nobody mistakes lexical overlap for real
#    semantic understanding.
# ---------------------------------------------------------------------------
def semantic_similarity(resume_text, jd_text) -> Dict:

    # Convert dicts to text
    if isinstance(resume_text, dict):
        resume_text = " ".join(str(v) for v in resume_text.values() if v)

    if isinstance(jd_text, dict):
        jd_text = " ".join(str(v) for v in jd_text.values() if v)

    # Convert everything else to string
    resume_text = str(resume_text or "").strip()
    jd_text = str(jd_text or "").strip()

    if not resume_text or not jd_text:
        return {
            "score": 0.0,
            "method": "no comparable text available"
        }

    # ... keep the rest of your function unchanged ...
    if not resume_text or not jd_text:
        return {"score": 0.0, "method": "no comparable text available"}

    model = _get_embedding_model()
    if model is not None:
        try:
            emb = model.encode([resume_text, jd_text])
            sim = cosine_similarity([emb[0]], [emb[1]])[0][0]
            return {
                "score": round(float(sim) * 100, 1),
                "method": "sentence-embeddings (all-MiniLM-L6-v2)",
            }
        except Exception:
            pass

    # Fallback: TF-IDF cosine similarity. This measures lexical/word-overlap
    # similarity, not true semantic meaning -- e.g. "built REST services" and
    # "developed APIs" will score lower here than they would with real
    # embeddings. We label it clearly rather than calling it "semantic".
    try:
        vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        tfidf = vectorizer.fit_transform([resume_text, jd_text])
        sim = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    except ValueError:
        # e.g. both texts are pure stopwords/empty after vectorization
        sim = 0.0
    return {
        "score": round(float(sim) * 100, 1),
        "method": "tfidf-cosine (lexical overlap fallback -- install sentence-transformers "
                  "with model weights available for true semantic scoring)",
    }


# ---------------------------------------------------------------------------
# 8. RESUME COMPLETENESS SCORE
# ---------------------------------------------------------------------------
REQUIRED_SECTIONS = ["summary", "education", "experience", "skills"]
BONUS_SECTIONS = ["projects", "certifications"]


def resume_completeness_score(sections: Dict[str, str], candidate: Dict) -> Dict:
    present_required = [s for s in REQUIRED_SECTIONS if sections.get(s)]
    missing_required = [s for s in REQUIRED_SECTIONS if not sections.get(s)]
    present_bonus = [s for s in BONUS_SECTIONS if sections.get(s)]
    missing_bonus = [s for s in BONUS_SECTIONS if not sections.get(s)]

    contact_fields = {
        "name": candidate["name"] != "Unknown Candidate",
        "email": candidate["email"] != "Not found",
        "phone": candidate["phone"] != "Not found",
    }
    contact_found = sum(contact_fields.values())

    section_score = 65 * (len(present_required) / len(REQUIRED_SECTIONS)) + \
                     15 * (len(present_bonus) / len(BONUS_SECTIONS))
    contact_score = 20 * (contact_found / len(contact_fields))
    total = round(section_score + contact_score, 1)

    feedback = []
    if missing_required:
        feedback.append(f"Add a clearly labeled section for: {', '.join(missing_required)}.")
    if missing_bonus:
        feedback.append(f"Consider adding: {', '.join(missing_bonus)} to strengthen the resume.")
    if not contact_fields["email"]:
        feedback.append("No email address was detected -- make sure it's near the top of the resume.")
    if not contact_fields["phone"]:
        feedback.append("No phone number was detected.")
    if not feedback:
        feedback.append("All key sections and contact details were detected.")

    return {
        "score": min(100.0, total),
        "present_sections": present_required + present_bonus,
        "missing_sections": missing_required + missing_bonus,
        "contact_complete": contact_found == len(contact_fields),
        "feedback": feedback,
    }


# ---------------------------------------------------------------------------
# 9. RESUME STRUCTURE SCORE (parseability / formatting)
# ---------------------------------------------------------------------------
BULLET_CHARS = ("-", "*", "•", "◦", "▪", "o ")


def resume_structure_score(resume_text: str, sections: Dict[str, str]) -> Dict:
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    word_count = len(resume_text.split())

    # 1. Header clarity: how many of the standard sections were detectable?
    header_ratio = len(sections) / len(SECTION_HEADERS) if SECTION_HEADERS else 0

    # 2. Bullet usage inside experience/projects (scannable achievements
    #    vs. dense paragraphs)
    body = (sections.get("experience", "") + "\n" + sections.get("projects", "")).split("\n")
    body_lines = [l.strip() for l in body if l.strip()]
    bullet_lines = [l for l in body_lines if l.startswith(BULLET_CHARS)]
    bullet_ratio = (len(bullet_lines) / len(body_lines)) if body_lines else 0

    # 3. Length check -- too short reads as thin, too long reads as unfocused
    if 250 <= word_count <= 1100:
        length_score = 100
    elif word_count < 250:
        length_score = round(max(30, 100 * (word_count / 250)), 1)
    else:
        length_score = round(max(50, 100 - (word_count - 1100) / 20), 1)

    header_score = round(header_ratio * 100, 1)
    bullet_score = round(bullet_ratio * 100, 1)

    total = round(0.45 * header_score + 0.30 * bullet_score + 0.25 * length_score, 1)

    feedback = []
    if header_score < 60:
        feedback.append("Use standard section headings (e.g. 'Experience', 'Education') so parsers can find them.")
    if bullet_score < 40:
        feedback.append("Use bullet points for experience and project entries instead of dense paragraphs.")
    if word_count < 250:
        feedback.append("The resume looks thin -- add more detail on responsibilities and outcomes.")
    if word_count > 1100:
        feedback.append("The resume is quite long -- tighten it to the most relevant, recent experience.")
    if not feedback:
        feedback.append("Formatting is clear and should parse reliably.")

    return {
        "score": total,
        "header_clarity": header_score,
        "bullet_usage": bullet_score,
        "length_score": length_score,
        "word_count": word_count,
        "feedback": feedback,
    }


# ---------------------------------------------------------------------------
# 10. PROJECT ANALYSIS / PROJECT QUALITY
# ---------------------------------------------------------------------------
ACTION_VERBS = [
    "built", "designed", "developed", "led", "implemented", "optimized", "automated",
    "architected", "launched", "created", "improved", "reduced", "increased", "migrated",
    "deployed", "engineered", "scaled", "streamlined", "integrated",
]

# Quantifiable-impact detector: looks for percentages, multipliers, or a
# number paired with a metric-ish word. Deliberately does NOT try to match
# calendar years/dates -- that's not a signal of project quality.
METRIC_RE = re.compile(
    r"\d+%|\b\d+x\b|\$\s?\d[\d,]*|\b\d[\d,]*\+?\s*(?:users|requests|records|transactions|"
    r"customers|downloads|rows|ms|seconds|hours|days|tickets|deployments)\b",
    re.IGNORECASE,
)


def analyze_projects(sections: Dict[str, str], jd_skills: List[str]) -> Dict:
    text = sections.get("projects", "")
    if not text.strip():
        return {
            "score": 0.0,
            "project_count_estimate": 0,
            "has_quantified_impact": False,
            "relevant_tech_mentioned": [],
            "action_verbs_used": [],
            "feedback": ["No projects section was detected -- add 1-3 projects relevant to the roles you're targeting."],
        }

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    # Rough project count: lines that look like a title (short, no bullet
    # prefix) approximate one entry per project.
    title_like = [l for l in lines if not l.startswith(BULLET_CHARS) and len(l) < 90]
    project_count = max(1, len(title_like))

    lowered = text.lower()
    relevant_tech = [s for s in jd_skills if s in lowered]
    verbs_used = [v for v in ACTION_VERBS if v in lowered]
    has_metrics = bool(METRIC_RE.search(text))

    tech_coverage = (len(relevant_tech) / len(jd_skills)) if jd_skills else (1.0 if relevant_tech else 0.5)
    verb_score = min(1.0, len(verbs_used) / 4)
    metric_score = 1.0 if has_metrics else 0.3

    score = round(100 * (0.45 * tech_coverage + 0.30 * metric_score + 0.25 * verb_score), 1)

    feedback = []
    if not has_metrics:
        feedback.append("Add measurable outcomes to your projects (e.g. '30% faster', '10k+ users').")
    if len(verbs_used) < 2:
        feedback.append("Lead project bullets with strong action verbs (built, designed, optimized...).")
    if jd_skills and not relevant_tech:
        feedback.append("None of the job description's key skills show up in your projects -- consider adding a relevant one.")
    if not feedback:
        feedback.append("Projects are well-aligned, quantified, and use strong action language.")

    return {
        "score": score,
        "project_count_estimate": project_count,
        "has_quantified_impact": has_metrics,
        "relevant_tech_mentioned": relevant_tech,
        "action_verbs_used": verbs_used,
        "feedback": feedback,
    }


# ---------------------------------------------------------------------------
# 11. EDUCATION REQUIREMENT EXTRACTION
#    FIX: JD bullet points usually don't end in a period, so splitting the
#    whole JD by sentence-final punctuation alone merges an entire bullet
#    list (and everything after it, up to the NEXT period anywhere in the
#    document) into one giant "sentence". That garbage blob was being
#    compared against the resume's education line, tanking the similarity
#    score even when the actual requirement matched fine. Splitting on
#    newlines FIRST (then sentence punctuation within each line) keeps each
#    bullet point self-contained.
# ---------------------------------------------------------------------------
# Sentinel returned by extract_education_requirement() when the JD simply
# doesn't mention an education requirement. Callers check for this and treat
# it the same way missing experience requirements are treated (0 years
# required -> full marks) rather than penalizing the candidate for not
# matching text that was never actually a requirement.
NO_EDU_REQUIREMENT_TEXT = "No specific education requirement stated in the job description."


import re

NO_EDU_REQUIREMENT = {
    "degree": None,
    "fields": [],
    "related": False
}

def extract_field(text):

    text = text.lower()

    mapping = {
        "computer science and engineering": "computer science",
        "computer science": "computer science",
        "information technology": "information technology",
        "data science": "data science",
        "statistics": "statistics",
        "mathematics": "mathematics",
        "economics": "economics",
        "electronics": "electronics",
        "mechanical": "mechanical",
        "civil": "civil",
        "electrical": "electrical"
    }

    for key, value in mapping.items():
        if key in text:
            return value

    return ""
def normalize_education(text):

    text = text.lower()

    if any(x in text for x in [
        "phd",
        "ph.d"
    ]):
        return "PhD"

    if any(x in text for x in [
        "master",
        "m.tech",
        "mtech",
        "mba",
        "msc",
        "m.s"
    ]):
        return "Master"

    if any(x in text for x in [
        "bachelor",
        "b.tech",
        "btech",
        "b.e",
        "be ",
        "bachelor of technology"
    ]):
        return "Bachelor"

    return "Unknown"

def extract_education_requirement(jd_text: str):

    lowered = jd_text.lower()

    # ---------------- Degree ---------------- #
    degree_patterns = {
        r"\b(ph\.?d|doctor of philosophy)\b": "PhD",
        r"\b(master(?:'s)?|masters|m\.?tech|mba|mca|m\.?sc)\b": "Master",
        r"\b(bachelor(?:'s)?|b\.?tech|b\.?e|bachelor of technology|bachelor of engineering)\b": "Bachelor",
        r"\b(diploma)\b": "Diploma"
    }

    degree = None
    # If no explicit degree is mentioned but the JD says "graduate",
# treat it as a Bachelor's requirement.
    if degree is None:
        if re.search(r"\b(graduate|undergraduate|graduation)\b", lowered):
            degree = "Bachelor"

    for pattern, label in degree_patterns.items():
        if re.search(pattern, lowered):
            degree = label
            break

    # ---------------- Fields ---------------- #
    field_patterns = {
        r"\bcomputer science\b": "Computer Science",
        r"\bcomputer science and engineering\b": "Computer Science",
        r"\bcse\b": "Computer Science",

        r"\bdata science\b": "Data Science",
        r"\bstatistics\b": "Statistics",
        r"\bmathematics\b": "Mathematics",
        r"\beconomics\b": "Economics",
        r"\binformation technology\b": "Information Technology",
        r"\bsoftware engineering\b": "Software Engineering",
        r"\bartificial intelligence\b": "Artificial Intelligence"
        
    }

    fields = []

    for pattern, label in field_patterns.items():
        if re.search(pattern, lowered):
            fields.append(label)

    fields = list(dict.fromkeys(fields))

    related = "related field" in lowered or "related discipline" in lowered

    return {
        "degree": degree,
        "fields": fields,
        "related": related
    }

# ---------------------------------------------------------------------------
# 12. SECTION-WISE ANALYSIS + RECOMMENDATIONS
# ---------------------------------------------------------------------------
import re

def normalize_education(text: str):
    text = text.lower()

    # ---------- Degree ----------
    degree = None

    if re.search(r"\b(ph\.?d|doctor of philosophy)\b", text):
        degree = "phd"

    elif re.search(r"\b(master|master's|masters|m\.?tech|mtech|mba|mca|m\.?sc|msc|m\.?e|me)\b", text):
        degree = "master"

    elif re.search(
        r"\b(bachelor|bachelor's|bachelors|b\.?tech|btech|b\.?e|be|b\.?sc|bsc|bca|bba|b\.?com|bcom)\b",
        text,
    ):
        degree = "bachelor"

    elif "diploma" in text:
        degree = "diploma"

    # ---------- Field ----------
    field_patterns = {
        r"\bcomputer science and engineering\b": "computer science",
        r"\bcomputer science\b": "computer science",
        r"\bcse\b": "computer science",

        r"\binformation technology\b": "information technology",
        r"\bsoftware engineering\b": "software engineering",

        r"\bdata science\b": "data science",
        r"\bartificial intelligence\b": "artificial intelligence",
        r"\bmachine learning\b": "machine learning",

        r"\bstatistics\b": "statistics",
        r"\bmathematics\b": "mathematics",
        r"\beconomics\b": "economics",

        r"\bmechanical\b": "mechanical engineering",
        r"\bcivil\b": "civil engineering",
        r"\belectrical\b": "electrical engineering",
        r"\belectronics\b": "electronics engineering",
    }

    fields = []

    for pattern, value in field_patterns.items():
        if re.search(pattern, text):
            fields.append(value)

    return {
        "degree": degree,
        "fields": list(dict.fromkeys(fields))
    }


# Education is considered a real match at a lower bar than "excellent" --
# genuinely relevant-but-not-word-for-word education text (very common,
# since resumes and JDs almost never phrase this identically) shouldn't be
# reported as "missing".
EDUCATION_MATCH_THRESHOLD = 55


def analyze_resume_sections(
    sections: Dict[str, str],
    jd_text: str,
    resume_text: str,
    resume_skills: Optional[List[str]] = None,
    jd_skills: Optional[List[str]] = None,
    skill_match: Optional[Dict] = None,
) -> Dict:
    """Section-by-section comparison used by the UI tabs.

    Technical/soft skill matching here now reuses the SAME semantic
    skill_match engine (exact + synonym + embedding-based matches) used for
    the headline score and the global skill chips, instead of a separate,
    narrower substring check against only the segmented "skills" section
    text. That narrower check was also the reason a genuinely-present
    skill could show as "missing" in this tab whenever section segmentation
    didn't perfectly capture the skills section -- matching now falls back
    to skills found anywhere in the resume, and is scored by meaning
    (exact / synonym / semantic), not literal keyword presence alone.
    """
    resume_skills = resume_skills if resume_skills is not None else find_skills(resume_text)
    jd_skills = jd_skills if jd_skills is not None else find_skills(jd_text)
    skill_match = skill_match if skill_match is not None else semantic_skill_match(resume_text, resume_skills, jd_skills)

    jd_skill_keys = jd_skills
    jd_technical = [s for s in jd_skill_keys if s not in SOFT_SKILLS]
    jd_soft = [s for s in jd_skill_keys if s in SOFT_SKILLS]

    analysis = {}

    # ---------- Education ----------
    resume_education = sections.get("education", "") or resume_text
    jd_education = extract_education_requirement(jd_text)

    if (
    jd_education == NO_EDU_REQUIREMENT_TEXT
    or (
        isinstance(jd_education, dict)
        and jd_education.get("degree") is None
        and not jd_education.get("fields")
    )
):
        analysis["education"] = {
            "resume": sections.get("education", ""),
            "job_requirement": jd_education,
            "matched": ["No specific education requirement to satisfy"],
            "missing": [],
            "similarity": 100.0,
            "feedback": "The job description doesn't state a specific education requirement.",
        }
    else:
        print("resume_education:", type(resume_education), resume_education)
        print("jd_education:", type(jd_education), jd_education)
        if isinstance(jd_education, dict):
            jd_education_text = " ".join(
                str(v) for v in jd_education.values() if v
            )
        else:
            jd_education_text = str(jd_education)

        # Extract resume education as structured information
    resume_info = extract_education_requirement(resume_education)

    score = 0

    # Degree match
    if (
        resume_info["degree"]
        and jd_education["degree"]
        and resume_info["degree"] == jd_education["degree"]
    ):
        score += 70

    # Field match
    if jd_education["fields"]:
        matched_fields = set(resume_info["fields"]) & set(jd_education["fields"])

        if matched_fields:
            score += 30 * len(matched_fields) / len(jd_education["fields"])

    score = round(score, 1)

    # ---- Build specific, text-level matched/missing entries instead of a
    #      single generic "Education requirement matched/missing" label, so
    #      the candidate can see exactly *what* in the JD their resume does
    #      or doesn't cover. ----
    DEGREE_RANK = {"Diploma": 0, "Bachelor": 1, "Master": 2, "PhD": 3}
    jd_degree = jd_education.get("degree")
    resume_degree = resume_info.get("degree")
    jd_fields = jd_education.get("fields") or []
    resume_fields = resume_info.get("fields") or []

    edu_matched, edu_missing = [], []

    if jd_degree:
        if resume_degree and (
            resume_degree == jd_degree
            or DEGREE_RANK.get(resume_degree, -1) > DEGREE_RANK.get(jd_degree, 99)
        ):
            edu_matched.append(
                f"Degree — JD requires a {jd_degree}'s degree; your resume shows a "
                f"{resume_degree}'s degree, which satisfies it."
            )
        else:
            resume_degree_text = f"{resume_degree}'s degree" if resume_degree else "no clearly stated degree"
            edu_missing.append(
                f"Degree — JD requires a {jd_degree}'s degree, but your resume shows "
                f"{resume_degree_text}."
            )

    if jd_fields:
        matched_fields = [f for f in jd_fields if f in resume_fields]
        missing_fields = [f for f in jd_fields if f not in resume_fields]
        for f in matched_fields:
            edu_matched.append(
                f"Field of study — JD asks for {f}, and this was found in your resume."
            )
        for f in missing_fields:
            edu_missing.append(
                f"Field of study — JD asks for {f}, but it wasn't detected in your "
                f"resume's education section."
            )

    if not jd_degree and not jd_fields:
        edu_matched = ["The job description doesn't state a specific degree or field of study."]
        edu_missing = []

    analysis["education"] = {
        "resume": sections.get("education", ""),
        "job_requirement": jd_education,
        "matched": edu_matched,
        "missing": edu_missing,
        "similarity": score,
        "feedback": (
            "Excellent education match."
            if score >= 80
            else "Education requirement is a reasonable match."
            if score >= EDUCATION_MATCH_THRESHOLD
            else "Education requirement is weakly matched."
        ),
    }
    # ---------- Technical Skills (meaning-based: exact + synonym + semantic) ----------
    tech_matches = skill_match["technical"]["matched"]
    tech_missing = skill_match["technical"]["missing"]
    match_label = {"exact": "exact match", "synonym": "synonym match", "semantic": "semantic match"}

    resume_technical_display = [s for s in resume_skills if s not in SOFT_SKILLS]

    analysis["technical_skills"] = {
        "resume": resume_technical_display,
        "job_requirement": jd_technical,
        "matched": [f"{m['skill']} ({match_label.get(m['match_type'], m['match_type'])})" for m in tech_matches],
        "missing": tech_missing,
        "feedback": (
            "Strong technical alignment with the job description." if not tech_missing else
            f"Add or highlight experience with: {', '.join(tech_missing[:5])}."
        ),
    }

    # ---------- Soft Skills (meaning-based) ----------
    soft_matches = skill_match["soft"]["matched"]
    soft_missing = skill_match["soft"]["missing"]
    resume_soft_display = [s for s in resume_skills if s in SOFT_SKILLS]

    analysis["soft_skills"] = {
        "resume": resume_soft_display,
        "job_requirement": jd_soft,
        "matched": [f"{m['skill']} ({match_label.get(m['match_type'], m['match_type'])})" for m in soft_matches],
        "missing": soft_missing,
        "feedback": (
            "Soft skills align well with the role." if not soft_missing else
            "Mention more teamwork, communication, and leadership examples."
        ),
    }

    # ---------- Experience ----------
    resume_exp = sections.get("experience", "")
    jd_exp_match = EXPERIENCE_RE.search(jd_text)
    required_years = float(jd_exp_match.group(1)) if jd_exp_match else 0.0

    resume_exp_match = EXPERIENCE_RE.search(resume_exp) or EXPERIENCE_RE.search(resume_text)
    candidate_years = float(resume_exp_match.group(1)) if resume_exp_match else 0.0

    exp_matched, exp_missing = [], []
    if required_years > 0 and candidate_years >= required_years:
        exp_matched.append(
            f"JD requires {required_years}+ years of experience; your resume shows "
            f"{candidate_years} years, which meets it."
        )
    elif required_years > 0:
        exp_missing.append(
            f"JD requires {required_years}+ years of experience, but your resume shows "
            f"{candidate_years if candidate_years else 'no clearly stated'} years."
        )

    analysis["experience"] = {
        "resume": resume_exp,
        "job_requirement": f"{required_years}+ years" if required_years else "Experience preferred",
        "matched": exp_matched,
        "missing": exp_missing,
        "feedback": (
            "Experience level satisfies the requirement." if not exp_missing else
            "Highlight internships, projects, or relevant work to strengthen experience alignment."
        ),
    }

    # ---------- Projects ----------
    resume_projects = sections.get("projects", "")
    project_keywords = [s for s in jd_technical if s in resume_projects.lower()]
    missing_project_keywords = [s for s in jd_technical if s not in resume_projects.lower()]

    analysis["projects"] = {
        "resume": resume_projects,
        "job_requirement": jd_technical,
        "matched": project_keywords,
        "missing": missing_project_keywords[:5],
        "feedback": (
            "Projects are relevant to the target role." if project_keywords else
            "Add projects demonstrating the required technologies."
        ),
    }

    # ---------- Certifications ----------
    resume_certs = sections.get("certifications", "")
    cert_keywords = ["aws", "azure", "gcp", "tensorflow", "docker", "kubernetes"]
    matched_certs = [c.upper() for c in cert_keywords if c in resume_certs.lower()]
    missing_certs = [c.upper() for c in cert_keywords if c in jd_text.lower() and c not in resume_certs.lower()]

    analysis["certifications"] = {
        "resume": resume_certs,
        "job_requirement": "Relevant technical certifications",
        "matched": matched_certs,
        "missing": missing_certs,
        "feedback": (
            "Relevant certifications strengthen the profile." if matched_certs else
            "Consider certifications related to the target role."
        ),
    }

    return analysis


def section_recommendations(section_analysis: Dict) -> List[Dict]:
    """Prioritized, section-wise recommendations: sections with missing
    items surface first so the person knows what to fix before anything
    that's already in good shape."""
    order = ["technical_skills", "experience", "education", "projects", "certifications", "soft_skills"]
    labels = {
        "technical_skills": "Technical Skills", "experience": "Experience", "education": "Education",
        "projects": "Projects", "certifications": "Certifications", "soft_skills": "Soft Skills",
    }
    recs = []
    for key in order:
        sect = section_analysis.get(key, {})
        recs.append({
            "section": labels[key],
            "needs_attention": bool(sect.get("missing")),
            "recommendation": sect.get("feedback", ""),
        })
    recs.sort(key=lambda r: not r["needs_attention"])
    return recs


# ---------------------------------------------------------------------------
# 13. EDUCATION / EXPERIENCE SCORING HELPERS
# ---------------------------------------------------------------------------
def get_education_rank(education_str: str) -> int:
    ed_lower = education_str.lower()
    if any(term in ed_lower for term in ["phd", "ph.d"]):
        return 4
    if any(term in ed_lower for term in ["master", "m.tech", "mtech", "m.sc", "mba"]):
        return 3
    if any(term in ed_lower for term in ["bachelor", "b.tech", "btech", "b.sc"]):
        return 2
    return 1


# ---------------------------------------------------------------------------
# 14. INTERVIEW READINESS + STRENGTHS/WEAKNESSES
# ---------------------------------------------------------------------------
def interview_readiness(job_match_score: float, missing_technical: List[str],
                         completeness_score: float, project_quality_score: float) -> Dict:
    if job_match_score >= 80 and completeness_score >= 75:
        level = "Ready"
        summary = "Your resume is well aligned with this role and complete enough to hold up under scrutiny."
    elif job_match_score >= 60:
        level = "Almost Ready"
        summary = "You're a reasonable fit -- closing a few skill or resume gaps would meaningfully help."
    elif job_match_score >= 40:
        level = "Needs Preparation"
        summary = "There are real gaps between your resume and this role. Prepare clear talking points for them."
    else:
        level = "Needs Significant Work"
        summary = "This resume and job description are not well aligned yet."

    talking_points = []
    if missing_technical:
        talking_points.append(
            f"Be ready to speak to: {', '.join(missing_technical[:4])} -- even brief hands-on familiarity helps."
        )
    if project_quality_score < 60:
        talking_points.append("Prepare 1-2 measurable outcomes for your projects (a number, %, or metric).")
    if completeness_score < 75:
        talking_points.append("Fill in missing resume sections before interview day -- interviewers may ask about gaps.")
    if not talking_points:
        talking_points.append("Focus your prep on the specifics of your strongest, most relevant projects.")

    return {"level": level, "summary": summary, "talking_points": talking_points}


def strengths_and_weaknesses(section_analysis: Dict, completeness: Dict, structure: Dict,
                              project_quality: Dict, skill_match: Dict) -> Dict:
    strengths, weaknesses = [], []

    if not skill_match["technical"]["missing"]:
        strengths.append("Strong technical skill coverage for this role.")
    elif len(skill_match["technical"]["missing"]) > max(2, len(skill_match["technical"]["matched"])):
        weaknesses.append("Significant technical skill gaps relative to the job description.")

    if completeness["score"] >= 85:
        strengths.append("Resume is complete -- all key sections and contact details are present.")
    elif completeness["missing_sections"]:
        weaknesses.append(f"Missing sections: {', '.join(completeness['missing_sections'])}.")

    if structure["score"] >= 80:
        strengths.append("Resume is well-structured and should parse cleanly through an ATS.")
    elif structure["score"] < 60:
        weaknesses.append("Formatting/structure could make this harder for an ATS to parse correctly.")

    if project_quality["has_quantified_impact"]:
        strengths.append("Projects include measurable, quantified impact.")
    else:
        weaknesses.append("Projects lack quantified outcomes (numbers, %, scale).")

    if section_analysis.get("experience", {}).get("missing"):
        weaknesses.append("Experience section falls short of the years the JD asks for.")
    elif section_analysis.get("experience", {}).get("matched"):
        strengths.append("Experience meets or exceeds what the job description asks for.")

    return {"strengths": strengths or ["No standout strengths detected yet."],
            "weaknesses": weaknesses or ["No major weaknesses detected."]}


# ---------------------------------------------------------------------------
# 15. FINAL RESULT + ORCHESTRATION
# ---------------------------------------------------------------------------
@dataclass
class AnalysisResult:
    candidate: Dict
    job_match_score: float          # headline: how well resume fits THIS job (skills/exp/edu/semantic)
    semantic_similarity: Dict       # {"score":.., "method":..}
    resume_completeness: Dict
    resume_structure: Dict
    project_analysis: Dict
    hiring_recommendation: str
    interview_readiness: Dict
    strengths_weaknesses: Dict
    skill_match: Dict
    resume_skills_by_category: Dict
    jd_skills_by_category: Dict
    sections: Dict = field(default_factory=dict)
    section_analysis: Dict = field(default_factory=dict)
    recommendations: List = field(default_factory=list)
    skills_score: float = 0.0
    education_score: float = 0.0
    experience_score: float = 0.0
    explanation: Dict = field(default_factory=dict)

    def to_dict(self) -> Dict:
        """Structured JSON for downstream modules (e.g. a recruiter API,
        a candidate-facing report, or another service in the pipeline)."""
        return asdict(self)


def analyze(resume_text: str, jd_text: str) -> AnalysisResult:
    candidate = extract_candidate_profile(resume_text)
    sections = segment_resume_sections(resume_text)

    resume_skills = find_skills(resume_text)
    jd_skills = find_skills(jd_text)
    resume_skills_by_category = categorize_skills(resume_skills)
    jd_skills_by_category = categorize_skills(jd_skills)

    # Semantic skill matching (exact + synonym + embedding-based) computed
    # ONCE and reused everywhere (global chips + section-wise tabs) so the
    # "meaning-based" scoring is consistent across the whole UI.
    skill_match = semantic_skill_match(resume_text, resume_skills, jd_skills)

    section_analysis = analyze_resume_sections(
        sections, jd_text, resume_text,
        resume_skills=resume_skills, jd_skills=jd_skills, skill_match=skill_match,
    )

    sim = semantic_similarity(resume_text, jd_text)

    completeness = resume_completeness_score(sections, candidate)
    structure = resume_structure_score(resume_text, sections)
    project_quality = analyze_projects(sections, jd_skills)

    # ---- Skills score: exact + synonym + semantic matches all count,
    #      weighted slightly by how they were matched ----
    match_weight = {"exact": 1.0, "synonym": 0.9, "semantic": 0.75}
    matched_weighted = sum(match_weight.get(m["match_type"], 0.5) for m in skill_match["matched"])
    skills_score = round(100 * matched_weighted / len(jd_skills), 1) if jd_skills else 100.0

    # ---- Education score (semantic similarity between candidate's actual
    #      education line and the JD's education requirement) ----
    # ------------------ Education Matching ------------------

    resume_education = sections.get("education", "") or resume_text
    jd_education = extract_education_requirement(jd_text)

    if jd_education == NO_EDU_REQUIREMENT_TEXT:
        education_score = 100.0
        education_feedback = "No education requirement specified."

    else:

        # ---------- Normalize Resume Degree ----------
        resume_degree = normalize_education(resume_education)

        # ---------- Extract Resume Field ----------
        resume_field = extract_field(resume_education)

        jd_degree = jd_education.get("degree", "")
        jd_fields = [f.lower() for f in jd_education.get("fields", [])]

        degree_score = 0
        field_score = 0
        semantic_score = 0

        # Degree Match (40%)
        # Ensure values are strings
        resume_degree = str(resume_degree or "")
        jd_degree = str(jd_degree or "")
        resume_field = str(resume_field or "")

        print("resume_degree:", resume_degree, type(resume_degree))
        print("jd_degree:", jd_degree, type(jd_degree))

        # Degree Match (40%)
        if resume_degree.lower() == jd_degree.lower():
            degree_score = 40

        # Field Match (40%)
        if resume_field.lower() in jd_fields:
            field_score = 40

        # Semantic Similarity (20%)
        # NOTE: use a separate variable here — do NOT overwrite `jd_text`,
        # since it's still needed below (e.g. experience-requirement
        # extraction). Overwriting it previously caused the Experience
        # score in the breakdown to silently show 100% (no years found)
        # even when the Experience tab correctly flagged it as missing.
        jd_edu_text = (
            f"{jd_degree} degree in "
            f"{', '.join(jd_education.get('fields', []))}"
        )

        edu = semantic_similarity(resume_education, jd_edu_text)

        semantic_score = (edu["score"] / 100.0) * 20

        education_score = round(
            degree_score +
            field_score +
            semantic_score,
            1
        )

        education_score = min(100.0, education_score)

        education_feedback = (
            "Education requirement matched."
            if education_score >= 80
            else "Education partially matches."
        )
        # ---- Experience score ----
        candidate_years = candidate["years_experience"] or 0.0
        jd_exp_match = EXPERIENCE_RE.search(jd_text)
        required_years = float(jd_exp_match.group(1)) if jd_exp_match else 0.0
        if required_years == 0.0:
            experience_score = 100.0
        elif candidate_years >= required_years:
            experience_score = 100.0
        else:
            experience_score = round((candidate_years / required_years) * 100, 1)

        # ---- Headline Job Match Score: skills + experience + education +
        #      real resume<->JD semantic similarity, weighted ----
        job_match_score = round(
            0.35 * skills_score +
            0.20 * experience_score +
            0.10 * education_score +
            0.20 * sim["score"] +
            0.15 * project_quality["score"],
            1
        )

        critical_missing = skill_match["technical"]["missing"]
        if job_match_score >= 85 and len(critical_missing) <= 1:
            recommendation = "Excellent Match — Prioritize for Interview"
        elif job_match_score >= 75:
            recommendation = "Strong Match — Shortlist for Interview"
        elif job_match_score >= 60:
            recommendation = "Good Match — Consider for Screening Round"
        elif job_match_score >= 45:
            recommendation = "Moderate Match — Needs Improvement"
        else:
            recommendation = "Weak Match — Not Recommended"

        readiness = interview_readiness(
            job_match_score, [s for s in critical_missing],
            completeness["score"], project_quality["score"],
        )
        sw = strengths_and_weaknesses(section_analysis, completeness, structure, project_quality, skill_match)
        recs = section_recommendations(section_analysis)

        explanation = {
            "skills_match": {
                "value": f"Skills Score: {skills_score}% "
                        f"({len(skill_match['matched'])} of {len(jd_skills)} required skills matched -- "
                        f"{sum(1 for m in skill_match['matched'] if m['match_type'] != 'exact')} of those via "
                        f"synonym/semantic matching, not exact text)",
            },
            "semantic_similarity": {
                "value": f"Resume-JD semantic similarity: {sim['score']}% (method: {sim['method']})",
            },
            "experience_comparison": {
                "value": f"Experience Score: {experience_score}% (Candidate has {candidate_years} yrs, JD requires {required_years} yrs)",
            },
            "education_match": {
                "value": f"Education Score: {education_score}% (Candidate: {candidate['highest_education']}, JD preferred: {jd_education})",
            },
            "project_quality": {
                "value": f"Project Quality Score: {project_quality['score']}%",
            },
            "overall_scoring": {
                "value": (f"Job Match Score breakdown: 35% Skills ({skills_score}%) + 20% Experience "
                        f"({experience_score}%) + 10% Education ({education_score}%) + 20% Semantic "
                        f"Similarity ({sim['score']}%) + 15% Project Quality ({project_quality['score']}%) "
                        f"= {job_match_score}%"),
            },
        }

    return AnalysisResult(
        candidate=candidate,
        job_match_score=job_match_score,
        semantic_similarity=sim,
        resume_completeness=completeness,
        resume_structure=structure,
        project_analysis=project_quality,
        hiring_recommendation=recommendation,
        interview_readiness=readiness,
        strengths_weaknesses=sw,
        skill_match=skill_match,
        resume_skills_by_category=resume_skills_by_category,
        jd_skills_by_category=jd_skills_by_category,
        sections=sections,
        section_analysis=section_analysis,
        recommendations=recs,
        skills_score=skills_score,
        education_score=education_score,
        experience_score=experience_score,
        explanation=explanation,
    )


# ---------------------------------------------------------------------------
# DEMO — run this file directly to see real output
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    SAMPLE_RESUME = """
    Ananya Verma
    ananya.verma@mail.com | +91 98765 43210

    Summary
    Backend engineer with 4.5 years of experience building scalable APIs
    and ML-driven data pipelines in fintech.

    Skills
    Python, SQL, FastAPI, Django, Machine Learning, NLP, TensorFlow,
    Power BI, Git, communication skills, mentoring junior engineers.

    Experience
    Backend Engineer II, Finlytics (2022 - Present)
    Built and maintained REST APIs serving 2M+ daily requests using FastAPI.
    Led a team of 3 engineers on the fraud-detection pipeline rewrite.

    Backend Engineer I, Finlytics (2021 - 2022)
    Migrated legacy SQL reporting jobs into a Power BI dashboard suite.

    Projects
    Fraud Detection Pipeline - built an NLP + Machine Learning pipeline
    to flag suspicious transaction narratives in real time, cutting manual
    review time by 40%.
    Internal API Gateway - designed a FastAPI-based gateway used by 6
    internal teams.

    Certifications
    AWS Certified Developer - Associate
    TensorFlow Developer Certificate

    Education
    B.Tech in Computer Science, Osmania University, 2021
    """

    SAMPLE_JD = """
    Senior Backend Engineer — Platform Team

    We're looking for an engineer with 4+ years building scalable APIs in
    Python. Strong experience with FastAPI or Django, SQL/NoSQL databases,
    and machine learning pipelines preferred. Must have hands-on experience
    with Docker, Kubernetes, and CI/CD pipelines on AWS or Azure. Familiarity
    with Spark for large-scale data processing is a plus. Strong
    communication skills required.
    """

    result = analyze(SAMPLE_RESUME, SAMPLE_JD)

    print(json.dumps({
        "candidate": result.candidate,
        "job_match_score": result.job_match_score,
        "semantic_similarity": result.semantic_similarity,
        "hiring_recommendation": result.hiring_recommendation,
        "interview_readiness": result.interview_readiness,
        "strengths_weaknesses": result.strengths_weaknesses,
        "resume_completeness": result.resume_completeness,
        "resume_structure": result.resume_structure,
        "project_analysis": result.project_analysis,
        "resume_skills_by_category": result.resume_skills_by_category,
        "skill_match": result.skill_match,
        "recommendations": result.recommendations,
        "explanation": result.explanation,
    }, indent=2))